import io
import logging
import threading

import docx
import pdfplumber
from celery import shared_task
from django.conf import settings
from django.db import close_old_connections
from django.utils import timezone

from apps.core.storage import download_file

from .models import ParsedResume, Resume
from .resume_parser import parse_resume_text

logger = logging.getLogger(__name__)


def _background_mode() -> str:
    """How to run post-upload work: "celery" (default), "thread", or eager.

    "thread" exists for deployments that run only the web process (no Celery
    worker): tasks queued with .delay() would sit in Redis forever, so we run
    them in a daemon thread inside the web process instead.
    """
    return str(getattr(settings, "BACKGROUND_TASKS_MODE", "celery")).lower()


def _run_in_thread(func, *args) -> None:
    def runner():
        try:
            func(*args)
        finally:
            close_old_connections()

    threading.Thread(target=runner, daemon=True).start()


def dispatch_resume_parse(resume_id: str) -> None:
    if getattr(settings, "CELERY_TASK_ALWAYS_EAGER", False):
        parse_resume_with_llm(str(resume_id))
        return

    if _background_mode() == "thread":
        _run_in_thread(parse_resume_with_llm, str(resume_id))
        return

    try:
        parse_resume_with_llm.delay(str(resume_id))
    except Exception:
        logger.warning("Celery unavailable; parsing resume %s inline.", resume_id)
        parse_resume_with_llm(str(resume_id))


def dispatch_resume_extraction(resume: "Resume", file_bytes: bytes | None = None) -> None:
    """Run text extraction + parsing for a freshly uploaded resume.

    Honors CELERY_TASK_ALWAYS_EAGER (synchronous, using in-memory bytes when
    available) and BACKGROUND_TASKS_MODE=thread (web-process background thread,
    for deployments without a Celery worker). Defaults to a Celery task.
    """
    if file_bytes is not None and getattr(settings, "CELERY_TASK_ALWAYS_EAGER", False):
        extract_resume_text_from_bytes(resume, file_bytes)
        return

    if _background_mode() == "thread":
        _run_in_thread(extract_resume_text, str(resume.id))
        return

    try:
        extract_resume_text.delay(str(resume.id))
    except Exception:
        logger.warning("Celery unavailable; extracting resume %s in a thread.", resume.id)
        _run_in_thread(extract_resume_text, str(resume.id))


def extract_text_from_bytes(file_bytes: bytes, mime_type: str) -> str:
    file_obj = io.BytesIO(file_bytes)

    if mime_type == "application/pdf":
        return extract_pdf_text(file_obj)

    if mime_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        doc = docx.Document(file_obj)
        return extract_docx_text(doc)

    logger.warning("Unsupported resume mime type: %s", mime_type)
    return ""


def extract_pdf_text(file_obj: io.BytesIO) -> str:
    """Extract text from a PDF using layout-aware extraction.

    Uses ``layout=True`` so pdfplumber respects the visual column/row
    structure of the page.  This prevents multi-column resumes from having
    their left and right column text interleaved.  For each page we also
    extract any tables separately and append them so skills/education stored
    inside table cells are not silently dropped.
    """
    pages: list[str] = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            parts: list[str] = []

            # Layout-aware extraction preserves visual column order.
            body = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
            if body and body.strip():
                parts.append(body.strip())

            # Extract tables so cell content is not dropped when it lives
            # inside cells that layout extraction misses.
            try:
                for table in page.extract_tables():
                    if not table:
                        continue
                    table_lines: list[str] = []
                    seen_cells: set[str] = set()
                    for row in table:
                        row_cells: list[str] = []
                        for cell in row:
                            cell_text = (cell or "").strip()
                            if cell_text and cell_text not in seen_cells:
                                seen_cells.add(cell_text)
                                row_cells.append(cell_text)
                        if row_cells:
                            table_lines.append("  ".join(row_cells))
                    if table_lines:
                        parts.append("\n".join(table_lines))
            except Exception as table_exc:  # noqa: BLE001
                logger.debug("PDF table extraction failed on page: %s", table_exc)

            if parts:
                pages.append("\n".join(parts))

    return "\n\n".join(pages)


def extract_docx_text(doc: docx.Document) -> str:
    """Extract text from a python-docx Document in reading order.

    Paragraphs are emitted first (body order), then each table cell on its
    own line so the LLM sees clean unambiguous content.  python-docx repeats
    the same Cell object for horizontally merged cells, so we deduplicate
    within each row to avoid redundant text that inflates the token count.
    """
    lines: list[str] = []

    def add_line(value: str) -> None:
        text = " ".join((value or "").split())
        if text:
            lines.append(text)

    # Section headers / footers (contact info often lives here)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            add_line(paragraph.text)
        for paragraph in section.footer.paragraphs:
            add_line(paragraph.text)

    # Body paragraphs
    for paragraph in doc.paragraphs:
        add_line(paragraph.text)

    # Tables — each unique cell on its own line, deduplicating merged cells
    for table in doc.tables:
        for row in table.rows:
            seen_in_row: set[str] = set()
            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                # python-docx repeats the same Cell object for merged cells
                if cell_text and cell_text not in seen_in_row:
                    seen_in_row.add(cell_text)
                    add_line(cell_text)

    return "\n".join(lines)


def extract_resume_text_from_bytes(resume: Resume, file_bytes: bytes) -> None:
    try:
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["status", "updated_at"])

        resume.raw_text = extract_text_from_bytes(file_bytes, resume.mime_type)
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["raw_text", "status", "updated_at"])

        dispatch_resume_parse(str(resume.id))
    except Exception as exc:
        logger.error("Failed to extract text for resume %s: %s", resume.id, exc)
        resume.status = Resume.Status.ERROR
        resume.save(update_fields=["status", "updated_at"])


@shared_task
def extract_resume_text(resume_id: str):
    try:
        resume = Resume.objects.get(id=resume_id)
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["status", "updated_at"])

        # Fetch the file from Supabase to memory
        file_bytes = download_file("resumes", resume.file_url)
        extracted_text = extract_text_from_bytes(file_bytes, resume.mime_type)

        resume.raw_text = extracted_text
        resume.status = Resume.Status.PROCESSING
        resume.save(update_fields=["raw_text", "status", "updated_at"])

        dispatch_resume_parse(str(resume.id))

    except Exception as e:
        logger.error(f"Failed to extract text for resume {resume_id}: {e}")
        try:
            resume = Resume.objects.get(id=resume_id)
            resume.status = Resume.Status.ERROR
            resume.save(update_fields=["status", "updated_at"])
        except Resume.DoesNotExist:
            pass


@shared_task(
    autoretry_for=(Exception,),
    retry_backoff=True,
    retry_jitter=True,
    retry_kwargs={"max_retries": 2},
)
def parse_resume_with_llm(resume_id: str):
    resume = Resume.objects.select_related("candidate", "application").get(id=resume_id)
    resume.status = Resume.Status.PROCESSING
    resume.save(update_fields=["status", "updated_at"])

    parsed_resume, _created = ParsedResume.objects.update_or_create(
        resume=resume,
        defaults={
            "candidate": resume.candidate,
            "application": resume.application,
            "status": ParsedResume.Status.PROCESSING,
        },
    )

    try:
        result = parse_resume_text(resume.raw_text)
    except Exception as exc:
        logger.error(f"Failed to parse resume {resume_id}: {exc}")
        parsed_resume.status = ParsedResume.Status.ERROR
        parsed_resume.validation_errors = [str(exc)]
        parsed_resume.parsed_at = timezone.now()
        parsed_resume.save(
            update_fields=["status", "validation_errors", "parsed_at", "updated_at"]
        )
        resume.status = Resume.Status.ERROR
        resume.save(update_fields=["status", "updated_at"])
        return

    parsed_resume.data = result.data
    parsed_resume.confidence = result.confidence
    parsed_resume.parser_model = result.model
    parsed_resume.validation_errors = result.validation_errors or []
    parsed_resume.token_usage = result.token_usage or {}
    parsed_resume.estimated_cost = result.estimated_cost or 0
    parsed_resume.status = ParsedResume.Status.COMPLETED
    parsed_resume.parsed_at = timezone.now()
    parsed_resume.save(
        update_fields=[
            "data",
            "confidence",
            "parser_model",
            "validation_errors",
            "token_usage",
            "estimated_cost",
            "status",
            "parsed_at",
            "updated_at",
        ]
    )

    resume.status = Resume.Status.COMPLETED
    resume.save(update_fields=["status", "updated_at"])

    try:
        from apps.ai_engine.embeddings import update_parsed_resume_embedding

        update_parsed_resume_embedding(parsed_resume, force=True)
    except Exception as exc:
        logger.warning(
            "Failed to generate embedding for parsed resume %s: %s",
            parsed_resume.id,
            exc,
        )

    if resume.application_id:
        try:
            from apps.ai_engine.ranking import score_application

            score_application(resume.application, force=True)
        except Exception as exc:
            logger.warning(
                "Failed to refresh score for application %s after parsing resume %s: %s",
                resume.application_id,
                resume.id,
                exc,
            )
